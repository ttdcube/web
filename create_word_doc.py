
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Create document
doc = Document()

# Set default font
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
doc.styles['Normal'].font.size = Pt(12)

# Title
title = doc.add_heading('KẾ HOẠCH DỰ ÁN: HỆ THỐNG ĐẶT LỊCH KHÁM BỆNH ONLINE (CLINICFLOW)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph()
subtitle.add_run('Ngày tạo: 29/05/2026').italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Table of Contents
doc.add_heading('Mục lục', level=1)
toc_items = [
    "1. Nghiên cứu nền tảng",
    "2. Phân tích yêu cầu hệ thống",
    "3. Thiết kế User Flow",
    "4. Thiết kế cơ sở dữ liệu chuẩn",
    "5. Đề xuất kiến trúc hệ thống tối ưu",
    "6. Giải thích lý do chọn kiến trúc",
    "7. Lộ trình thực hiện (Sprint Planning)"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: Nghiên cứu nền tảng
doc.add_heading('1. Nghiên cứu nền tảng', level=1)
p = doc.add_paragraph()
p.add_run('Dựa trên nghiên cứu các hệ thống đặt lịch khám hiện đại, dashboard quản lý bệnh viện và giao diện health-tech, chúng tôi rút ra các điểm chính:')

research_points = [
    'Thiết kế ưu tiên di động (Mobile-first): 80% người dùng truy cập trang healthcare qua điện thoại di động',
    'Quyền truy cập dựa trên vai trò (RBAC - Role-Based Access Control): 3 vai trò chính: Bệnh nhân, Bác sĩ, Quản trị viên',
    'Thiết kế giao diện tin cậy: Màu xanh dương/xanh lá cây tạo cảm giác yên tâm, không rườm rà',
    'Luồng đặt lịch rõ ràng: Nhiều bước với chỉ báo tiến trình, yêu cầu thông tin tối thiểu ban đầu',
    'RESTful API chuẩn: Sử dụng JWT để xác thực, phân lớp rõ ràng'
]
for point in research_points:
    doc.add_paragraph(point, style='List Bullet')

# Section 2: Phân tích yêu cầu
doc.add_heading('2. Phân tích yêu cầu hệ thống', level=1)

doc.add_heading('Yêu cầu chức năng', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Module'
hdr_cells[1].text = 'Yêu cầu chi tiết'

functional_reqs = [
    ("Xác thực người dùng", "Đăng ký/Đăng nhập bằng JWT; Mật khẩu mã hóa BCrypt; Phân quyền dựa trên vai trò"),
    ("Hồ sơ bệnh nhân", "Xem/Cập nhật thông tin cá nhân; Lưu thông tin y tế (nhóm máu, dị ứng, tiền sử bệnh)"),
    ("Đặt lịch khám", "Tìm kiếm bác sĩ theo chuyên khoa; Xem lịch trống thời gian thực; Xác nhận lịch; Nhắc lại tự động"),
    ("Quản lý lịch bác sĩ", "Xem lịch tuần; Đánh dấu khoảng thời gian bận; Xác nhận/từ chối yêu cầu đặt lịch; Xem danh sách bệnh nhân"),
    ("Dashboard quản trị", "Quản lý người dùng; Tổng quan về lịch khám; Thống kê, báo cáo"),
    ("Tránh trùng lịch", "Hoạt động nguyên tử (Atomic Operations) để khóa slot trong quá trình đặt lịch")
]
for module, req in functional_reqs:
    row_cells = table.add_row().cells
    row_cells[0].text = module
    row_cells[1].text = req

doc.add_heading('Yêu cầu phi chức năng', level=2)
non_func_reqs = [
    'Hiệu suất: Thời gian phản hồi < 2 giây cho các thao tác đặt lịch',
    'An ninh: Dữ liệu được mã hóa, tuân thủ các quy định bảo mật',
    'Khả năng mở rộng: Hỗ trợ hơn 1000 lượt đặt lịch đồng thời',
    'Khả dụng: 99.9% thời gian hoạt động',
    'Tính dễ sử dụng: Độ tương phản cao, tuân thủ tiêu chuẩn WCAG'
]
for req in non_func_reqs:
    doc.add_paragraph(req, style='List Bullet')

# Section 3: User Flow
doc.add_page_break()
doc.add_heading('3. Thiết kế User Flow', level=1)

doc.add_heading('3.1 Luồng của Bệnh nhân', level=2)
doc.add_paragraph('Trang chủ → (Đăng nhập tùy chọn) → Tìm kiếm bác sĩ → Chọn bác sĩ → Xem lịch trống → Chọn ngày/giờ → Điền thông tin đặt lịch → Xác nhận → Dashboard cá nhân')

doc.add_heading('3.2 Luồng của Bác sĩ', level=2)
doc.add_paragraph('Đăng nhập → Dashboard → Xem lịch hôm nay → Đánh dấu bận → Xem xét lịch hẹn → Xem hồ sơ bệnh nhân')

doc.add_heading('3.3 Luồng của Quản trị viên', level=2)
doc.add_paragraph('Đăng nhập → Tổng quan Dashboard → Quản lý bác sĩ/bệnh nhân → Xem thống kê → Cấu hình hệ thống')

# Section 4: Database
doc.add_page_break()
doc.add_heading('4. Thiết kế cơ sở dữ liệu chuẩn (3NF)', level=1)

doc.add_heading('4.1 Bảng Người dùng (Users)', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Cột'
hdr[1].text = 'Kiểu dữ liệu'
hdr[2].text = 'Mô tả'

users_data = [
    ('user_id', 'INT (PK, Auto Increment)', 'ID người dùng'),
    ('username', 'VARCHAR(50) UNIQUE', 'Tên đăng nhập'),
    ('email', 'VARCHAR(100) UNIQUE', 'Email'),
    ('password_hash', 'VARCHAR(255)', 'Mật khẩu mã hóa'),
    ('role', "ENUM('PATIENT','DOCTOR','ADMIN')", 'Vai trò'),
    ('created_at', 'DATETIME', 'Thời gian tạo')
]
for row in users_data:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

doc.add_heading('4.2 Bảng Bệnh nhân (Patients)', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Cột'
hdr[1].text = 'Kiểu dữ liệu'
hdr[2].text = 'Mô tả'

patients_data = [
    ('patient_id', 'INT (PK, Auto Increment)', 'ID bệnh nhân'),
    ('user_id', 'INT (FK, UNIQUE)', 'Liên kết với bảng Users'),
    ('full_name', 'VARCHAR(100)', 'Họ tên đầy đủ'),
    ('phone', 'VARCHAR(20)', 'Số điện thoại'),
    ('birth_date', 'DATE', 'Ngày sinh'),
    ('gender', "ENUM('M','F','OTHER')", 'Giới tính'),
    ('blood_type', 'VARCHAR(5)', 'Nhóm máu'),
    ('allergies', 'TEXT', 'Dị ứng'),
    ('medical_history', 'TEXT', 'Tiền sử bệnh')
]
for row in patients_data:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

doc.add_heading('4.3 Bảng Lịch hẹn (Appointments)', level=2)
table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Cột'
hdr[1].text = 'Kiểu dữ liệu'
hdr[2].text = 'Mô tả'

appointments_data = [
    ('appointment_id', 'INT (PK, Auto Increment)', 'ID lịch hẹn'),
    ('patient_id', 'INT (FK)', 'Liên kết với bảng Patients'),
    ('doctor_id', 'INT (FK)', 'Liên kết với bảng Doctors'),
    ('slot_id', 'INT (FK)', 'Liên kết với bảng Doctor Availability'),
    ('reason', 'TEXT', 'Lý do khám'),
    ('status', "ENUM('PENDING','CONFIRMED','COMPLETED','CANCELLED')", 'Trạng thái'),
    ('created_at', 'DATETIME', 'Thời gian tạo')
]
for row in appointments_data:
    cells = table.add_row().cells
    cells[0].text = row[0]
    cells[1].text = row[1]
    cells[2].text = row[2]

# Section 5: Architecture
doc.add_page_break()
doc.add_heading('5. Đề xuất kiến trúc hệ thống tối ưu', level=1)
doc.add_paragraph('Hệ thống được thiết kế theo kiến trúc 3 lớp:')

architecture = """
┌─────────────────────────────────────────┐
│ Presentation Layer (Giao diện)          │
│ • React SPA (Responsive)                │
└─────────────────────────────────────────┘
                    ↓
┌─────────────────────────────────────────┐
│ Application Layer (Backend)             │
│ • Node.js + Express + RESTful API       │
└─────────────────────────────────────────┘
                    ↓
┌─────────────────────────────────────────┐
│ Data Layer                              │
│ • MySQL (Dữ liệu quan hệ) + Redis       │
└─────────────────────────────────────────┘
"""
doc.add_paragraph(architecture)

# Section 6: Why this architecture
doc.add_heading('6. Giải thích lý do chọn kiến trúc', level=1)
reasons = [
    'React + Express + MySQL: MERN-adjacent stack, đã được chứng minh cho các hệ thống healthcare booking, dễ phát triển và bảo trì',
    'Redis: Ngăn chặn race condition trong quá trình đặt lịch, cache lịch trống của bác sĩ để truy cập nhanh hơn, giảm tải cơ sở dữ liệu',
    'Tách biệt các lớp rõ ràng: Dễ phát triển, kiểm thử, mở rộng'
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

# Section 7: Roadmap
doc.add_heading('7. Lộ trình thực hiện (Sprint Planning)', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Sprint'
hdr[1].text = 'Kết quả cần đạt được'

sprints = [
    ('Sprint 1', 'Cấu trúc dự án; Định nghĩa bảng CSDL; Xác thực JWT; Giao diện đăng ký/đăng nhập'),
    ('Sprint 2', 'CRUD hồ sơ bệnh nhân và bác sĩ; Lưu thông tin y tế bệnh nhân; Quản lý chuyên khoa bác sĩ'),
    ('Sprint 3', 'Quản lý thời gian trống; Luồng đặt lịch với kiểm tra trùng lặp'),
    ('Sprint 4', 'Portal bệnh nhân; Lịch bác sĩ; Tổng quan admin với thống kê'),
    ('Sprint 5', 'Tinh chỉnh UI/UX; Nhắc nhở tự động; Tích hợp GitHub; Triển khai lên hosting')
]
for sprint, deliverable in sprints:
    cells = table.add_row().cells
    cells[0].text = sprint
    cells[1].text = deliverable

# Save document
doc.save('KEHOACH_DU_AN_CLINICFLOW.docx')
print("Đã tạo file KEHOACH_DU_AN_CLINICFLOW.docx thành công!")
